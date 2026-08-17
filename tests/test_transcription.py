"""Tests de la chaîne de transcription.

Le moteur Whisper lui-même n'est pas testé ici (il faut le modèle) : on utilise
le moteur factice pour vérifier tout ce qui vient après — identification des
intervenants, relevé de décisions, et écriture des livrables.
"""
import json
import sys
import tempfile
from itertools import permutations
from pathlib import Path

import numpy as np

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from audiotool import diarize, documents, minutes  # noqa: E402
from audiotool.pipeline import Settings, process_file, run_transcription  # noqa: E402
from audiotool.transcribe import Segment, TranscribeSettings  # noqa: E402

FAILS = []


def check(name, cond, detail=""):
    print(f"{'✓' if cond else '✗'} {name} {detail}")
    if not cond:
        FAILS.append(name)


# --------------------------------------------------------------------------- #
#  Échéances et porteurs d'action
# --------------------------------------------------------------------------- #

CAS = [
    ("Je m'occupe de prévenir le client avant vendredi.", "avant vendredi"),
    ("Il faut que Marie prépare le support pour le comité de direction.", None),
    ("On livre pour le 12 mars.", "pour le 12 mars"),
    ("Je reviens vers toi d'ici la fin de semaine.", "d'ici la fin de semaine"),
    ("Réponse sous 48 heures.", "sous 48 heures"),
    ("On se voit la semaine prochaine.", "la semaine prochaine"),
    ("Le budget est de 42 000 euros.", None),
    ("On en reparle le 15/03.", "15/03"),
    ("Il faut boucler avant le 30.", "avant le 30"),
    ("On a validé les 3 lots hier.", None),
]
ok = sum((minutes._deadline(t) == w) if w else (minutes._deadline(t) is None) for t, w in CAS)
check("détection des échéances", ok == len(CAS), f"({ok}/{len(CAS)})")

check("porteur : première personne → le locuteur",
      minutes._owner(dict(text="Je m'occupe du dossier.", speaker="Intervenant 2")) == "Intervenant 2")
check("porteur : prénom cité dans la phrase",
      minutes._owner(dict(text="Il faut que Marie prépare le support.", speaker="Intervenant 1")) == "Marie")

# --------------------------------------------------------------------------- #
#  Relevé de décisions
# --------------------------------------------------------------------------- #

segs = [
    Segment(0, 5, "Bonjour, on commence par le budget.", "Intervenant 1"),
    Segment(5, 11, "On est à 42 000 euros sur le trimestre.", "Intervenant 2"),
    Segment(11, 17, "D'accord, on décide de reporter la livraison au 15 mars.", "Intervenant 1"),
    Segment(17, 23, "Je m'occupe de prévenir le client avant vendredi.", "Intervenant 2"),
    Segment(23, 28, "Est-ce qu'on a une visibilité sur les recrutements ?", "Intervenant 1"),
    Segment(28, 34, "Il faut que Marie prépare le support.", "Intervenant 2"),
]
rel = minutes.extract(segs, 34.0, diarize.speaking_time(segs))
check("décision repérée", len(rel["decisions"]) == 1, f"({len(rel['decisions'])})")
check("actions repérées", len(rel["actions"]) == 2, f"({len(rel['actions'])})")
check("question ouverte repérée", len(rel["questions"]) == 1)
check("montant repéré", len(rel["chiffres"]) == 1)
check("échéance de l'action", rel["actions"][0]["echeance"] == "avant vendredi",
      f"({rel['actions'][0]['echeance']!r})")
check("porteur de l'action", rel["actions"][1]["porteur"] == "Marie",
      f"({rel['actions'][1]['porteur']!r})")
parts = rel["stats"]["temps_parole"]
check("temps de parole calculé", set(parts) == {"Intervenant 1", "Intervenant 2"}
      and parts["Intervenant 1"]["part"].endswith("%"))

# --------------------------------------------------------------------------- #
#  Attribution des segments aux tours de parole
# --------------------------------------------------------------------------- #

turns = [diarize.Turn(0, 10, "Intervenant 1"), diarize.Turn(10, 20, "Intervenant 2")]
test = [Segment(1, 4, "a"), Segment(12, 15, "b"), Segment(9, 11, "c"), Segment(25, 27, "d")]
diarize.assign_speakers(test, turns)
check("attribution par recouvrement",
      [s.speaker for s in test[:2]] == ["Intervenant 1", "Intervenant 2"])
check("segment à cheval → tour majoritaire", test[2].speaker in ("Intervenant 1", "Intervenant 2"))
check("segment hors des tours → tour le plus proche", test[3].speaker == "Intervenant 2")

# --------------------------------------------------------------------------- #
#  Identification des intervenants sur dialogues de référence
# --------------------------------------------------------------------------- #


def precision(turns_pred, truth):
    labs = sorted({t.speaker for t in turns_pred})
    ref = sorted({s["speaker"] for s in truth})
    best = 0.0
    for perm in permutations(labs, min(len(ref), len(labs))):
        mp = {perm[i]: ref[i] for i in range(len(perm))}
        good = tot = 0
        for sg in truth:
            for x in np.arange(sg["start"] + 0.4, sg["end"] - 0.4, 0.1):
                p = next((t.speaker for t in turns_pred if t.start <= x <= t.end), None)
                tot += 1
                good += (p is not None and mp.get(p) == sg["speaker"])
        best = max(best, good / max(tot, 1))
    return best


sys.path.insert(0, str(Path(__file__).parent))
import make_dialogue  # noqa: E402

with tempfile.TemporaryDirectory() as tmp:
    for n in (1, 2, 3):
        src = f"{tmp}/dlg{n}.m4a"
        make_dialogue.main(src, n_speakers=n, n_turns=12)
        res = process_file(src, tmp, Settings(exports=["transcribe"]))
        wav = res["outputs"][0]["path"]
        truth = json.load(open(src.replace(".m4a", "_verite.json")))
        pred = diarize._light_diarize(wav)
        k = len({t.speaker for t in pred})
        pr = precision(pred, truth)
        check(f"intervenants ({n} voix) : nombre détecté", k == n, f"(détecté {k})")
        check(f"intervenants ({n} voix) : attribution", pr > 0.85, f"({pr * 100:.0f} %)")

    # --------------------------------------------------------------------- #
    #  Livrables
    # --------------------------------------------------------------------- #
    ts = TranscribeSettings(enabled=True, engine="mock", formats=["docx", "txt", "srt", "json"])
    out = run_transcription(wav, tmp, ts, duration=res["report"]["duree_s"], titre="Réunion test")
    kinds = {o["kind"]: Path(o["path"]) for o in out["outputs"]}
    check("quatre livrables écrits", len(kinds) == 4, f"({sorted(kinds)})")
    for kind, path in kinds.items():
        check(f"livrable non vide : {kind}", path.exists() and path.stat().st_size > 500,
              f"({path.stat().st_size if path.exists() else 0} octets)")

    srt = kinds["sous_titres"].read_text()
    check("format SRT valide", srt.startswith("1\n") and " --> " in srt.splitlines()[1])
    data = json.loads(kinds["donnees_json"].read_text())
    check("JSON exploitable", {"meta", "segments", "releve"} <= set(data)
          and len(data["segments"]) > 3 and "start" in data["segments"][0])

    from docx import Document
    doc = Document(str(kinds["compte_rendu"]))
    titres = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    check("document Word structuré", "Verbatim" in titres and any("Relevé" in t for t in titres),
          f"({titres[:4]})")
    check("tableaux présents dans le Word", len(doc.tables) >= 2, f"({len(doc.tables)})")

# --------------------------------------------------------------------------- #
#  Vérification du moteur (autotest)
# --------------------------------------------------------------------------- #

from audiotool import selftest, transcribe as _tr  # noqa: E402

CAS_WER = [
    ("bonjour à tous merci", "bonjour à tous merci", 0.0),
    ("bonjour à tous", "bonjour a tous", 0.0),            # accents ignorés
    ("bonjour à tous merci", "bonjour tous merci", 0.25),  # une omission
    ("bonjour à tous", "bonjour à tous les amis", 2 / 3),  # deux insertions
    ("le client vendredi", "le client mardi", 1 / 3),      # une substitution
]
bons = sum(abs(selftest.word_error_rate(r, h)["taux"] - a) < 0.01 for r, h, a in CAS_WER)
check("taux d'erreur sur les mots", bons == len(CAS_WER), f"({bons}/{len(CAS_WER)})")

d = selftest.word_error_rate("le client vendredi", "le client mardi")
check("types d'erreur distingués",
      (d["substitutions"], d["omissions"], d["insertions"]) == (1, 0, 0), f"({d})")

# branche « précision et vitesse » : on simule une voix de synthèse et un moteur
_vrai_transcribe = _tr.transcribe


def _faux_synthese(text, out):
    selftest.make_silence(Path(out), seconds=12.0)     # peu importe le contenu
    return True


def _faux_moteur(wav, settings, duration=0.0, progress=None):
    import time as _t
    _t.sleep(0.05)
    if "silence" in str(wav):
        return [], dict(langue="fr", modele=settings.model)
    texte = selftest.REFERENCE.replace("vendredi", "mardi")   # une erreur volontaire
    return [_tr.Segment(0.0, 12.0, texte)], dict(langue="fr", modele=settings.model)


selftest.synthesize = _faux_synthese
_tr.transcribe = _faux_moteur
try:
    rap = selftest.run("small", log=lambda m: None, engine="mock")
finally:
    _tr.transcribe = _vrai_transcribe

check("rapport : précision mesurée", rap["precision"] is not None
      and 0 < rap["precision"]["taux_erreur"] < 20,
      f"({rap['precision']['taux_erreur']} %)")
check("rapport : une substitution repérée", rap["precision"]["substitutions"] == 1)
check("rapport : vitesse mesurée", rap["vitesse"]["facteur"] > 1,
      f"({rap['vitesse']['facteur']}x)")
check("rapport : durée de calcul pour une heure",
      rap["vitesse"]["minutes_calcul_par_heure_audio"] > 0)
check("rapport : aucune hallucination sur le silence",
      rap["hallucinations"]["mots_inventes"] == 0)
check("rapport : verdict favorable", rap["verdict"]["niveau"] == "ok",
      f"({rap['verdict']['resume']})")

# et sans voix de synthèse disponible
selftest.synthesize = lambda text, out: False
_tr.transcribe = _faux_moteur
try:
    rap2 = selftest.run("small", log=lambda m: None, engine="mock")
finally:
    _tr.transcribe = _vrai_transcribe
check("rapport : dégradation propre sans voix système",
      rap2["precision"] is None and rap2["verdict"]["niveau"] == "partiel")


# --------------------------------------------------------------------------- #
#  Horodatage
# --------------------------------------------------------------------------- #
check("horodatage hh:mm:ss", documents.hhmmss(3725) == "01:02:05")
check("horodatage sous-titres", documents.hhmmss(3725.5, ms=True) == "01:02:05,500")

print("\nRÉSULTAT :", "OK" if not FAILS else f"ÉCHECS → {FAILS}")
sys.exit(1 if FAILS else 0)
