"""Mise en forme des livrables : texte, sous-titres, JSON et document Word."""
from __future__ import annotations

import json
from pathlib import Path


def hhmmss(t: float, ms: bool = False) -> str:
    t = max(0.0, float(t))
    h, rem = divmod(int(t), 3600)
    m, s = divmod(rem, 60)
    if ms:
        return f"{h:02d}:{m:02d}:{s:02d},{int((t % 1) * 1000):03d}"
    return f"{h:02d}:{m:02d}:{s:02d}"


def _turns(segments):
    """Regroupe les segments consécutifs d'un même intervenant en tours de parole."""
    out = []
    for s in segments:
        if out and out[-1]["speaker"] == s.speaker and s.start - out[-1]["end"] < 3.0:
            out[-1]["text"] += " " + s.text
            out[-1]["end"] = s.end
        else:
            out.append(dict(speaker=s.speaker, start=s.start, end=s.end, text=s.text))
    return out


# --------------------------------------------------------------------------- #
#  Formats texte
# --------------------------------------------------------------------------- #


def write_txt(path: str, segments, meta: dict) -> str:
    lines = [meta.get("titre", "Transcription"), "=" * len(meta.get("titre", "Transcription")), ""]
    for k, v in meta.items():
        if k != "titre" and v:
            lines.append(f"{k} : {v}")
    lines.append("")
    for t in _turns(segments):
        who = f"{t['speaker']} " if t["speaker"] else ""
        lines.append(f"[{hhmmss(t['start'])}] {who}".rstrip())
        lines.append(t["text"].strip())
        lines.append("")
    Path(path).write_text("\n".join(lines), encoding="utf-8")
    return path


def write_srt(path: str, segments) -> str:
    out = []
    for i, s in enumerate(segments, 1):
        who = f"{s.speaker} : " if s.speaker else ""
        out += [str(i), f"{hhmmss(s.start, True)} --> {hhmmss(s.end, True)}",
                f"{who}{s.text.strip()}", ""]
    Path(path).write_text("\n".join(out), encoding="utf-8")
    return path


def write_json(path: str, segments, meta: dict, releve: dict | None = None) -> str:
    data = dict(meta=meta,
                segments=[s.to_dict() for s in segments],
                releve=releve or {})
    Path(path).write_text(json.dumps(data, indent=1, ensure_ascii=False), encoding="utf-8")
    return path


# --------------------------------------------------------------------------- #
#  Document Word
# --------------------------------------------------------------------------- #


def write_docx(path: str, segments, meta: dict, releve: dict,
               synthese: str | None = None) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    grey = RGBColor(0x60, 0x60, 0x5C)

    def small(text, italic=True):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic = italic
        r.font.size = Pt(9)
        r.font.color.rgb = grey
        return p

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.autofit = False
        for i, h in enumerate(headers):
            run = t.rows[0].cells[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9.5)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                run = cells[i].paragraphs[0].add_run("" if v is None else str(v))
                run.font.size = Pt(9.5)
        if widths:
            # largeurs doubles obligatoires : la grille du tableau ET chaque
            # cellule, sinon Word et LibreOffice répartissent tout également
            for i, w in enumerate(widths[:len(t.columns)]):
                t.columns[i].width = Inches(w)
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        doc.add_paragraph()
        return t

    # ---- en-tête ---------------------------------------------------------- #
    title = doc.add_heading(meta.get("titre", "Compte rendu de réunion"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info = " · ".join(f"{k} : {v}" for k, v in meta.items() if k != "titre" and v)
    small(info)

    stats = releve.get("stats", {})
    if stats.get("temps_parole"):
        doc.add_heading("Intervenants", level=1)
        table(["Intervenant", "Temps de parole", "Part"],
              [[k, v["duree"], v["part"]] for k, v in stats["temps_parole"].items()],
              widths=[2.4, 1.6, 1.0])
        small("Renommez les intervenants par Rechercher/Remplacer une fois le document ouvert.")

    # ---- synthèse rédigée (si un modèle local l'a produite) ---------------- #
    if synthese:
        doc.add_heading("Synthèse", level=1)
        for para in synthese.split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("##"):
                doc.add_heading(para.lstrip("# ").strip(), level=2)
            elif para.startswith(("- ", "* ")):
                doc.add_paragraph(para[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(para)

    # ---- relevé par règles ------------------------------------------------ #
    doc.add_heading("Relevé automatique", level=1)
    small("Repérage par mots-clés à relire : ce qui suit signale des passages à "
          "vérifier dans le verbatim, ce n'est pas une synthèse validée.", italic=True)

    if releve.get("decisions"):
        doc.add_heading("Décisions annoncées", level=2)
        table(["Heure", "Intervenant", "Passage"],
              [[d["temps"], d.get("intervenant") or "", d["texte"]] for d in releve["decisions"]],
              widths=[0.75, 1.25, 4.5])
    if releve.get("actions"):
        doc.add_heading("Actions évoquées", level=2)
        table(["Heure", "Porteur", "Échéance", "Passage"],
              [[a["temps"], a.get("porteur") or "", a.get("echeance") or "", a["texte"]]
               for a in releve["actions"]], widths=[0.75, 1.1, 1.25, 3.4])
    if releve.get("questions"):
        doc.add_heading("Questions restées ouvertes", level=2)
        for q in releve["questions"]:
            doc.add_paragraph(f"{q['temps']} — {q['texte']}", style="List Bullet")
        doc.add_paragraph()
    if releve.get("chiffres"):
        doc.add_heading("Chiffres cités", level=2)
        for c in releve["chiffres"]:
            doc.add_paragraph(f"{c['temps']} — {c['texte']}", style="List Bullet")
        doc.add_paragraph()
    if releve.get("sujets"):
        doc.add_heading("Déroulé", level=2)
        table(["Plage", "Mots saillants"],
              [[f"{s['debut']} – {s['fin']}", ", ".join(s["mots_cles"])] for s in releve["sujets"]],
              widths=[1.6, 4.9])

    # ---- verbatim --------------------------------------------------------- #
    doc.add_page_break()
    doc.add_heading("Verbatim", level=1)
    for t in _turns(segments):
        p = doc.add_paragraph()
        head = p.add_run(f"[{hhmmss(t['start'])}] " + (f"{t['speaker']}" if t["speaker"] else ""))
        head.bold = True
        head.font.size = Pt(9.5)
        head.font.color.rgb = RGBColor(0x2F, 0x6F, 0x5E)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        body = doc.add_paragraph(t["text"].strip())
        body.paragraph_format.space_after = Pt(9)

    doc.save(path)
    return path
